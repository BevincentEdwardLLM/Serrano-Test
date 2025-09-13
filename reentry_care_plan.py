import streamlit as st
import os
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"D:\serrano_react\service_account.json"
import shutil
from io import BytesIO
import pandas as pd
from sqlalchemy import create_engine
from google.cloud import bigquery
import pymysql
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
import traceback
import re

load_dotenv()

# ✅ BigQuery client
try:
    client = bigquery.Client()
except Exception as e:
    print(f"Could not initialize BigQuery client: {e}")
    client = None

# ✅ UI → Actual column names mapping
CANON_MAP = {
    # identifiers
    "Medical ID Number": "Medical ID Number",
    "Medi-Cal ID Number": "Medical ID Number",
    "medical_id_number": "Medical ID Number",
    "youth_name": "Name of the youth",
    "Name of the youth": "Name of the youth",
    "Name": "Name of the youth",
    "Medical ID": "Medical ID Number",

    # dates / appointments
    "actual_release_date": "Actual release date",
    "scheduled_appointments": "Scheduled Appointments",
    "Release Date": "Actual release date",
    "Appointments": "Scheduled Appointments",
    "court_dates": "Court dates",
    "Court Dates": "Court dates",

    # social/economic
    "income_and_benefits": "Income and benefits",
    "Income": "Income and benefits",
    "food_and_clothing": "Food & Clothing",
    "Food & Clothing": "Food & Clothing",
    "identification_documents": "Identification documents",
    "ID Docs": "Identification documents",
    "life_skills": "Life skills",
    "Life Skills": "Life skills",
    "family_and_children": "Family and children",
    "Family": "Family and children",
    "service_referrals": "Service referrals",
    "Service Referrals": "Service referrals",
    "home_modifications": "Home Modifications",
    "Home Modifications": "Home Modifications",
    "durable_medical_equipment": "Durable Medical Equipment",
    "Durable Equipment": "Durable Medical Equipment",
    "Screenings": "Screenings",
    "housing": "Housing",
    "Housing": "Housing",
    "employment": "Employment",
    "Employment": "Employment",
    "transportation": "Transportation",
    "Transportation": "Transportation",
    "Treatment History": "Treatment History",
    "Treatment History (mental health, physical health, substance use)": "Treatment History",

    # ✅ extra 9 fields from screenshot
    "Race/Ethnicity": "Race/Ethnicity",
    "Residential Address": "Residential Address",
    "Telephone": "Telephone",
    "Medi-Cal health plan assigned": "Medi-Cal health plan assigned",
    "Health Screenings": "Health Screenings",
    "Health Assessments": "Health Assessments",
    "Chronic Conditions": "Chronic Conditions",
    "Prescribed Medications": "Prescribed Medications",
    "Primary physician contacts": "Primary physician contacts",
    "Clinical Assessments": "Clinical Assessments",
    "Emergency contacts": "Emergency contacts",
    "case_notes": "Case Notes"
}

# Master list defining the exact display order of fields in the documents.
# This list is used to ensure consistency with the UI's logical flow.
DISPLAY_ORDER_REENTRY = [
    "Name of the youth",
    "Medical ID Number", # Moved Medical ID to a prominent position
    "Race/Ethnicity",
    "Telephone",
    "Residential Address",
    "Emergency contacts",
    "Identification documents",
    "Case Notes",
    "Actual release date",
    "Court dates",
    "Medi-Cal health plan assigned",
    "Health Screenings",
    "Health Assessments",
    "Chronic Conditions",
    "Prescribed Medications",
    "Clinical Assessments",
    "Screenings",
    "Primary physician contacts",
    "Durable Medical Equipment",
    "Treatment History",
    "Scheduled Appointments",
    "Housing",
    "Food & Clothing",
    "Transportation",
    "Income and benefits",
    "Home Modifications",
    "Employment",
    "Life skills",
    "Family and children",
    "Service referrals",
]

def set_table_borders(table, color_rgb=(0, 0, 0)):
    """Apply borders to a table manually (works even without Word styles)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color_rgb))
        borders.append(border)

    tblPr.append(borders)

def set_cell_border(cell, color_rgb=(0, 0, 0)):
    """Set the borders of a single cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color_rgb))
        borders.append(border)
    tcPr.append(borders)

def _set_run_font(run, name="Century Gothic", size_pt=None, color_rgb=None):
    """
    Force a run's font (including East Asia paths) to a specific font.
    """
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts or OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if rPr.rFonts is None:
        rPr.append(rFonts)

    run.font.name = name

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)

def force_document_font(doc, name="Century Gothic"):
    """
    Apply the desired font to:
    - Normal style (document default)
    - All existing paragraphs/runs
    - All existing tables (headers + cells)
    """
    base = doc.styles['Normal']
    base.font.name = name
    rPr = base._element.get_or_add_rPr()
    rFonts = rPr.rFonts or OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if rPr.rFonts is None:
        rPr.append(rFonts)

    for p in doc.paragraphs:
        for run in p.runs:
            _set_run_font(run, name=name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _set_run_font(run, name=name)

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Rename known variants to canonical column names."""
    if df is None or df.empty:
        return df
    rename_map = {k: v for k, v in CANON_MAP.items() if k in df.columns}
    return df.rename(columns=rename_map)

def normalize_selected_fields(selected_fields):
    """Map UI labels to canonical where needed (e.g., Medi-Cal -> Medical)."""
    normalized = []
    for field in selected_fields:
        # Strip the suffix like "(CM)" and then normalize
        clean_key = field.split(" (")[0].strip()
        normalized_key = CANON_MAP.get(clean_key, clean_key)
        normalized.append(normalized_key)
    return normalized

def get_case_notes(sql_dict, bq_dict, dict_representation):
    """Fetch Case Notes with fallback SQL → BQ → Excel."""
    possible_keys = ["Case Notes", "case_notes", "casenotes"]
    for key in possible_keys:
        if sql_dict.get(key):
            return sql_dict[key]
        if bq_dict.get(key):
            return bq_dict[key]
        if dict_representation.get(key):
            return dict_representation[key]
    return "No case notes available."

def parse_candidate_name(candidate_name):
    """
    Parses a string like "John Doe (Medical ID: 1234567890) - Address - Phone"
    to extract the name and ID. Returns a tuple (name, medical_id).
    """
    match = re.match(r"^(.*) \(Medical ID: (\d+)\).*$", candidate_name)
    if match:
        name = match.group(1).strip()
        medical_id = match.group(2).strip()
        return name, medical_id
    return candidate_name, None

def get_candidates_by_name(person_input: str):
    """
    Searches Excel, SQL, and BigQuery for all people with the same name.
    Returns a de-duplicated list of (name, medical_id).
    """
    candidates = []

    # Strip the ID if present from the input string
    person_name, _ = parse_candidate_name(person_input)
    if not person_name:
        person_name = person_input

    # Excel
    try:
        file_data = pd.read_excel("ExcelFiles/reentry5.xlsx")
        file_data = normalize_columns(file_data)
        if "Name of the youth" in file_data.columns:
            matches = file_data[
                file_data["Name of the youth"].astype(str).str.strip().str.lower()
                == person_name.strip().lower()
            ]
            for _, row in matches.iterrows():
                mid = str(row.get("Medical ID Number") or "").strip()
                if mid:
                    candidates.append((row["Name of the youth"], mid))
    except Exception as e:
        print("Excel search error:", e)

    # SQL
    try:
        sql_df = read_cloud_sql(person_name, medical_id=None)
        sql_df = normalize_columns(sql_df)
        for _, row in sql_df.iterrows():
            mid = str(row.get("Medical ID Number") or "").strip()
            name = row.get("Name of the youth")
            if mid and name:
                candidates.append((name, mid))
    except Exception as e:
        print("SQL search error:", e)

    # BigQuery
    try:
        bq_df = read_bigquery(person_name, medical_id=None)
        bq_df = normalize_columns(bq_df)
        for _, row in bq_df.iterrows():
            mid = str(row.get("Medical ID Number") or "").strip()
            name = row.get("Name of the youth")
            if mid and name:
                candidates.append((name, mid))
    except Exception as e:
        print("BigQuery search error:", e)

    unique = {}
    for name, mid in candidates:
        if mid not in unique:
            unique[mid] = name
    return [(name, mid) for mid, name in unique.items()]


def generate_reentry_care_plan(selected_fields, candidate_name, app_option):
    """
    Fetches data from Excel, Cloud SQL, and BigQuery,
    merges it based on selected fields, and returns a BytesIO Word document.
    """
    try:
        person_input, medical_id = parse_candidate_name(candidate_name)
        
        if not person_input:
            person_input = candidate_name.split('(')[0].strip()
            medical_id = None
        
        file_data = pd.read_excel("ExcelFiles/reentry5.xlsx")
        file_data = normalize_columns(file_data)
        if medical_id and "Medical ID Number" in file_data.columns:
            person_row = file_data[file_data["Medical ID Number"].astype(str) == str(medical_id)]
        else:
            person_row = file_data[file_data.get("Name of the youth", pd.Series(dtype=str)) == person_input]
        dict_representation = person_row.to_dict(orient="records")[0] if not person_row.empty else {}

        sql_record = read_cloud_sql(person_input, medical_id)
        sql_record = normalize_columns(sql_record)
        sql_dict = (sql_record.to_dict(orient="records")[0]
                    if isinstance(sql_record, pd.DataFrame) and not sql_record.empty else {})

        bq_record = read_bigquery(person_input, medical_id)
        bq_record = normalize_columns(bq_record)
        bq_dict = (bq_record.to_dict(orient="records")[0]
                    if isinstance(bq_record, pd.DataFrame) and not bq_record.empty else {})

        merged_dict = {}
        merged_dict.update(dict_representation)
        merged_dict.update(sql_dict)
        merged_dict.update(bq_dict)
        merged_dict.pop("id", None)
        
        # FIX: Explicitly add the medical ID back into the dictionary
        if medical_id:
            merged_dict["Medical ID Number"] = medical_id
        
        normalized_merged_dict = {k.strip(): v for k, v in merged_dict.items()}

        print("\n✅ FINAL MERGED DATA:", normalized_merged_dict)

        template_path = "data/Template.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found at: {template_path}")
        
        doc = Document(template_path)
        
        # Add the main title for the table section
        title_paragraph = doc.add_paragraph(f"{person_input}'s Reentry Care Plan")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.runs[0]
        title_run.bold = True
        _set_run_font(title_run, name="Century Gothic", size_pt=16, color_rgb=(0, 0, 0))

        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=2)
        
        try:
            table.style = "Table Grid"
        except KeyError:
            set_table_borders(table)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        
        selected_canonical_keys = normalize_selected_fields(selected_fields)

        for canonical_key in DISPLAY_ORDER_REENTRY:
            row_cells = table.add_row().cells
            row_cells[0].text = str(canonical_key)
            
            value = normalized_merged_dict.get(canonical_key, "Data Not Available")

            if canonical_key in selected_canonical_keys:
                row_cells[1].text = "" if pd.isna(value) else str(value)
                if pd.isna(value) or str(value).strip() == "" or str(value).lower() == "nan":
                    row_cells[1].text = "Data Not Available"
            else:
                row_cells[1].text = "Not Selected"

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    except Exception as e:
        print("❌ Error in generate_reentry_care_plan:", str(e))
        return None

def generate_data_validation_report(selected_fields, candidate_name, app_option):
    """
    Generates a Word document listing ALL fields, indicating if data is present
    and applying a colored border to the value cell.
    """
    try:
        person_input, medical_id = parse_candidate_name(candidate_name)
        file_data = pd.read_excel("ExcelFiles/reentry5.xlsx")
        file_data = normalize_columns(file_data)

        if medical_id and "Medical ID Number" in file_data.columns:
            person_row = file_data[file_data["Medical ID Number"].astype(str) == str(medical_id)]
        else:
            person_row = file_data[file_data.get("Name of the youth", pd.Series(dtype=str)) == person_input]
        dict_representation = person_row.to_dict(orient="records")[0] if not person_row.empty else {}

        sql_record = read_cloud_sql(person_input, medical_id)
        sql_record = normalize_columns(sql_record)

        bq_record = read_bigquery(person_input, medical_id)
        bq_record = normalize_columns(bq_record)

        sql_dict = (sql_record.to_dict(orient="records")[0]
                    if isinstance(sql_record, pd.DataFrame) and not sql_record.empty else {})
        bq_dict = (bq_record.to_dict(orient="records")[0]
                    if isinstance(bq_record, pd.DataFrame) and not bq_record.empty else {})

        merged_dict = {}
        merged_dict.update(dict_representation)
        merged_dict.update(sql_dict)
        merged_dict.update(bq_dict)
        merged_dict.pop("id", None)

        # FIX: Explicitly add the medical ID back into the dictionary
        if medical_id:
            merged_dict["Medical ID Number"] = medical_id

        # Load the existing template document
        template_path = "data/Template.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found at: {template_path}")
        doc = Document(template_path)

        # Add the report title
        title_paragraph = doc.add_paragraph(f"Data Validation Report for {person_input}")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.runs[0]
        title_run.bold = True
        _set_run_font(title_run, name="Century Gothic", size_pt=16, color_rgb=(0, 0, 0))

        doc.add_paragraph("")

        # Create a table with 'Field' and 'Status' headers
        table = doc.add_table(rows=1, cols=2)
        try:
            table.style = "Table Grid"
        except KeyError:
            set_table_borders(table)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Status"
        
        all_canonical_fields = list(dict.fromkeys(CANON_MAP.values()).keys())
        
        for canonical_key in DISPLAY_ORDER_REENTRY:
            row_cells = table.add_row().cells
            
            ui_label = canonical_key
            row_cells[0].text = str(ui_label)

            value = merged_dict.get(canonical_key)
            
            is_missing = pd.isna(value) or (isinstance(value, str) and value.strip() == "") or value is None

            if canonical_key == "Case Notes":
                notes_value = get_case_notes(sql_dict, bq_dict, dict_representation)
                is_missing = pd.isna(notes_value) or str(notes_value).strip() == "" or str(notes_value).lower() in ["nan", "none", "no case notes available."]
            
            if is_missing:
                row_cells[1].text = "Data Not Available"
                set_cell_border(row_cells[1], color_rgb=(255, 0, 0))
            else:
                row_cells[1].text = "Data Available"
                set_cell_border(row_cells[1], color_rgb=(0, 128, 0))
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    except Exception as e:
        print("❌ Error in generate_data_validation_report:", traceback.format_exc())
        st.error(f"Failed to generate data validation report: {e}")
        return None

def read_cloud_sql(person_input, medical_id=None):
    user = os.environ.get("CLOUD_SQL_USER")
    password = os.environ.get("CLOUD_SQL_PASSWORD")
    host = os.environ.get("CLOUD_SQL_HOST")
    database = "serrano"

    if not all([user, password, host]):
        print("SQL connection details missing from environment variables.")
        return pd.DataFrame()

    connection_url = f"mysql+pymysql://{user}:{password}@{host}/{database}"
    engine = create_engine(connection_url)
    if medical_id:
        query = f"SELECT * FROM SocialEconomicLogistics_backup WHERE medical_id_number='{medical_id}'"
    else:
        query = f"SELECT * FROM SocialEconomicLogistics_backup WHERE youth_name='{person_input}'"
    try:
        df = pd.read_sql(query, engine)
        return df
    except Exception as e:
        print(f"Error reading from Cloud SQL: {e}")
        return pd.DataFrame()

def read_bigquery(person_input, medical_id=None):
    if not client:
        return pd.DataFrame()

    if medical_id:
        query = """
            SELECT *
            FROM `genai-poc-424806.SerranoAdvisorsBQ.scalablefeaturesforBQ`
            WHERE medical_id_number = @mid
        """
        job_config = bigquery.QueryJobConfig(
            query_parameters=[bigquery.ScalarQueryParameter("mid", "STRING", str(medical_id))]
        )
    else:
        query = """
            SELECT *
            FROM `genai-poc-424806.SerranoAdvisorsBQ.scalablefeaturesforBQ`
            WHERE youth_name = @name
        """
        job_config = bigquery.QueryJobConfig(
            query_parameters=[bigquery.ScalarQueryParameter("name", "STRING", person_input)]
        )
    try:
        df = client.query(query, job_config=job_config).to_dataframe()
        return df
    except Exception as e:
        print(f"Error reading from BigQuery: {e}")
        return pd.DataFrame()

DB_CONFIG = {
    "host": os.environ.get("CLOUD_SQL_HOST", "34.44.69.178"),
    "user": os.environ.get("CLOUD_SQL_USER", "root"),
    "password": os.environ.get("CLOUD_SQL_PASSWORD", "SQLsql$123"),
    "database": "serrano",
}
try:
    conn = pymysql.connect(**DB_CONFIG)
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM SocialEconomicLogistics_backup LIMIT 3")
    rows = cursor.fetchall()
    print("✅ rows:", rows)
    cursor.close()
    conn.close()
except Exception as e:
    print("DB smoke test failed:", e)